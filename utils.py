from openpyxl import Workbook
import json
from docx import Document

# ---------- EXCEL ----------
def save_to_excel(data, filename="scraped_data.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "Scraped Data"

    # Dynamic headers based on scraper output
    ws.append(["Title", "Content", "Link"])

    for item in data:
        ws.append([
            item.get("title") or item.get("tag", ""),
            item.get("content") or item.get("text", ""),
            item.get("link", "")
        ])

    wb.save(filename)
    return filename


# ---------- JSON ----------
def save_to_json(data, filename="scraped_data.json"):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    return filename


# ---------- WORD ----------
def save_to_word(data, summary="", filename="scraped_data.docx"):
    doc = Document()

    # Title
    doc.add_heading('Smart Scraper Report', 0)

    # Summary Section
    if summary:
        doc.add_heading('AI Summary', level=1)
        doc.add_paragraph(summary)

    # Data Section
    doc.add_heading('Extracted Content', level=1)

    for item in data:
        title = item.get("title") or item.get("tag", "")
        content = item.get("content") or item.get("text", "")
        link = item.get("link", "")

        # Title
        doc.add_heading(title, level=2)

        # Content
        doc.add_paragraph(content)

        # Link
        if link and link.startswith("http"):
            p = doc.add_paragraph("Source: ")
            p.add_run(link).italic = True

    doc.save(filename)
    return filename
    